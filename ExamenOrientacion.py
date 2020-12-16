import tkinter
from tkinter import filedialog
from tkinter import ttk
import openpyxl
import matplotlib
import matplotlib.pyplot as plt
from PIL import ImageTk,Image
import matplotlib.patches as mpatches
import os
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor

global img

def browseFiles():
    filename = filedialog.askopenfilename(initialdir = "/", title = "Selecciona el archivo con las respuestas.", filetypes = (("Archivos excel","*.xlsx"),("all files","*.*")))
    direccion.configure(text=filename)
    documentoExcel=openpyxl.load_workbook(filename)
    hojaActiva = documentoExcel.active
    nombres=[]
    for i in range(2, hojaActiva.max_row+1):
        nombres.append(hojaActiva.cell(row = i, column = 4).value)
    nombres=sorted(nombres)
    nombre_alumno["values"] = nombres



def calcularResultados(hojaActiva, inicio, fila, max):
    suma=0
    for i in range(inicio, max+1, 10):
        suma+=hojaActiva.cell(row = fila, column = i).value
    return int(suma)

def buscarAlumno():
    filename=direccion.cget("text")
    documentoExcel=openpyxl.load_workbook(filename)
    hojaActiva = documentoExcel.active
    i=1
    while hojaActiva.cell(row = i, column = 2).value != nombre_alumno.get():
        i+=1
        if i>hojaActiva.max_row:
            break
    if i>hojaActiva.max_row:
        find="Alumno no encontrado"
    else:
        find=hojaActiva.cell(row = i, column = 2).value
    fila = i-1
    valores_inter=[]
    valores_apt=[]
    servicio_social=calcularResultados(hojaActiva, 28, fila, 87)
    valores_inter.append(servicio_social)
    ejecutivo_persuasiva=calcularResultados(hojaActiva, 29, fila, 87)
    valores_inter.append(ejecutivo_persuasiva)
    verbal=calcularResultados(hojaActiva, 30, fila, 87)
    valores_inter.append(verbal)
    artistico_plastica=calcularResultados(hojaActiva, 31, fila, 87)
    valores_inter.append(artistico_plastica)
    musical=calcularResultados(hojaActiva, 32, fila, 87)
    valores_inter.append(musical)
    organizacion=calcularResultados(hojaActiva, 33, fila, 87)
    valores_inter.append(organizacion)
    cientifica=calcularResultados(hojaActiva, 34, fila, 87)
    valores_inter.append(cientifica)
    calculo=calcularResultados(hojaActiva, 35, fila, 87)
    valores_inter.append(calculo)
    mecanico_constructiva=calcularResultados(hojaActiva, 36, fila, 87)
    valores_inter.append(mecanico_constructiva)
    aire_libre=calcularResultados(hojaActiva, 37, fila, 87)
    valores_inter.append(aire_libre)

    changeBg(ss_intereses, servicio_social)
    changeBg(ep_intereses, ejecutivo_persuasiva)
    changeBg(v_intereses, verbal)
    changeBg(ap_intereses, artistico_plastica)
    changeBg(m_intereses, musical)
    changeBg(o_intereses, organizacion)
    changeBg(cient_intereses, cientifica)
    changeBg(calc_intereses, calculo)
    changeBg(mc_intereses, mecanico_constructiva)
    changeBg(al_intereses, aire_libre)

    servicio_social=calcularResultados(hojaActiva, 87, fila, 146)
    valores_apt.append(servicio_social)
    ejecutivo_persuasiva=calcularResultados(hojaActiva, 88, fila, 146)
    valores_apt.append(ejecutivo_persuasiva)
    verbal=calcularResultados(hojaActiva, 89, fila, 146)
    valores_apt.append(verbal)
    artistico_plastica=calcularResultados(hojaActiva, 90, fila, 146)
    valores_apt.append(artistico_plastica)
    musical=calcularResultados(hojaActiva, 91, fila, 146)
    valores_apt.append(musical)
    organizacion=calcularResultados(hojaActiva, 92, fila, 146)
    valores_apt.append(organizacion)
    cientifica=calcularResultados(hojaActiva, 93, fila, 146)
    valores_apt.append(cientifica)
    calculo=calcularResultados(hojaActiva, 94, fila, 146)
    valores_apt.append(calculo)
    mecanico_constructiva=calcularResultados(hojaActiva, 95, fila, 146)
    valores_apt.append(mecanico_constructiva)
    aire_libre=calcularResultados(hojaActiva, 96, fila, 146)
    valores_apt.append(aire_libre)

    changeBg(ss_aptitudes, servicio_social)
    changeBg(ep_aptitudes, ejecutivo_persuasiva)
    changeBg(v_aptitudes, verbal)
    changeBg(ap_aptitudes, artistico_plastica)
    changeBg(m_aptitudes, musical)
    changeBg(o_aptitudes, organizacion)
    changeBg(cient_aptitudes, cientifica)
    changeBg(calc_aptitudes, calculo)
    changeBg(mc_aptitudes, mecanico_constructiva)
    changeBg(dm_aptitudes, aire_libre)

    inter_apt=["I-SS", "A-SS", "I-EP", "A-EP", "I-V", "A-V", "I-AP", "A-AP", "I-MS", "A-MS", "I-OG", "A-OG", "I-CT", "A-CT", "I-CL", "A-CL", "I-MC", "A-MC", "I-AL", "A-DM"]
    valores_intapt=[*sum(zip(valores_inter,valores_apt),())]

    graficar(inter_apt, valores_intapt)

def changeBg(label, valor):
    texto=label.cget("text")
    saltoLinea=texto.find("\n")
    saltoLinea+=1
    texto=texto[0:saltoLinea]
    if valor>16:
        label.configure(text =texto+str(valor)+"/24", bg="green")
    elif valor>8:
        label.configure(text =texto+str(valor)+"/24", bg="orange")
    else:
        label.configure(text =texto+str(valor)+"/24", bg="red")

def graficar(inter_apt, valores_intapt):
    colores=["blue", "red"]
    fig, ax = plt.subplots(figsize=(10, 3.8))
    #Colocamos una etiqueta en el eje Y
    ax.set_ylabel('Puntaje')
    ax.set_xlabel('Intereses y Aptitudes')
    ax.set_ylim([0, 25])
    blue_patch = mpatches.Patch(color='blue', label='Intereses')
    red_patch = mpatches.Patch(color='red', label='Aptitudes')
    ax.legend(handles=[blue_patch, red_patch])
    #Colocamos una etiqueta en el eje X
    ax.set_title(nombre_alumno.get())
    #Creamos la grafica de barras utilizando 'paises' como eje X y 'ventas' como eje y.
    plt.bar(inter_apt, valores_intapt, align = "center", width=1, color=colores)
    plt.savefig(nombre_alumno.get()+'_grafica.jpg')
    img = ImageTk.PhotoImage(Image.open(nombre_alumno.get()+"_grafica.jpg"))
    grafica.configure(image=img)
    grafica.image = img

def generarDocx():
    filename=direccion.cget("text")
    documentoExcel=openpyxl.load_workbook(filename)
    hojaActiva = documentoExcel.active
    i=1
    while hojaActiva.cell(row = i, column = 4).value != nombre_alumno.get():
        i+=1
        if i>hojaActiva.max_row:
            break
    if i>hojaActiva.max_row:
        find="Alumno no encontrado"
    else:
        find=hojaActiva.cell(row = i, column = 4).value
    fila = i

    nombrepdf=nombre_alumno.get()
    listanombrepdf=nombrepdf.split(' ')

    nombrepdf=""

    for elemento in listanombrepdf:
        nombrepdf+=elemento
    nombrepdf+="_reporte.docx"

    document = Document()
    document.add_heading("Resultados: "+nombre_alumno.get()+"\nEdad; "+str(int(hojaActiva.cell(row = i, column = 3).value))+" años. Género; "+hojaActiva.cell(row=i, column=2).value+". Localidad; "+hojaActiva.cell(row=i, column=148).value, 1)
    p1 = document.add_paragraph()
    p1.add_run('Árbol de la Vida:').bold = True

    for x in range(0, 7):
        p2 = document.add_paragraph()
        p2.add_run(hojaActiva.cell(row=1, column=5+x).value).italic = True
        p3 = document.add_paragraph(hojaActiva.cell(row=i, column=5+x).value)
    p4 = document.add_paragraph()
    p4.add_run("Razones para estudiar una carrera universitaria:").bold = True
    respuestasRazones=hojaActiva.cell(row=i, column=12).value.split('\n')
    for x in range(0, 4):
        p5 = document.add_paragraph(respuestasRazones[x])
    p6 = document.add_paragraph()
    p6.add_run('De niño deseaba ser:').bold = True
    p7 = document.add_paragraph()
    p7.add_run(hojaActiva.cell(row=i, column=13).value).italic = True
    document.add_page_break()
    p8 = document.add_paragraph()
    p8.add_run("Eligió las siguientes personalidades:").bold = True
    for x in range(0, 4):
        p9 = document.add_paragraph(hojaActiva.cell(row=i, column=14+x).value)
    p10 = document.add_paragraph()
    p10.add_run("Preguntas Finales:").bold = True

    for x in range(0, 10):
        p11 = document.add_paragraph()
        p11.add_run(hojaActiva.cell(row=1, column=18+x).value[:-24]).italic = True
        valorSeparado=hojaActiva.cell(row=i,column=18+x).value.split('\n')
        p11 = document.add_paragraph(valorSeparado[0]+", "+valorSeparado[1])
    table = document.add_table(rows=12, cols=2)

    table.cell(0,0).text="-----RESPUESTAS INVENTARIO-----"
    table.cell(1,0).text="INTERESES:"
    table.cell(1,1).text="APTITUDES:"

    row = table.rows[2]
    printInventario(ss_intereses, ss_aptitudes, row)
    row = table.rows[3]
    printInventario(ep_intereses, ep_aptitudes, row)
    row = table.rows[4]
    printInventario(v_intereses, v_aptitudes, row)
    row = table.rows[5]
    printInventario(ap_intereses, ap_aptitudes, row)
    row = table.rows[6]
    printInventario(m_intereses, m_aptitudes, row)
    row = table.rows[7]
    printInventario(o_intereses, o_aptitudes, row)
    row = table.rows[8]
    printInventario(cient_intereses, cient_aptitudes, row)
    row = table.rows[9]
    printInventario(calc_intereses, calc_aptitudes, row)
    row = table.rows[10]
    printInventario(mc_intereses, mc_aptitudes, row)
    row = table.rows[11]
    printInventario(al_intereses, dm_aptitudes, row)
    p14 = document.add_paragraph()
    p14.alignment = 1
    r14 = p14.add_run()
    r14.add_picture(nombre_alumno.get()+'_grafica.jpg', width=Inches(7.5))

    document.save(nombrepdf)
    os.remove(nombre_alumno.get()+'_grafica.jpg')
    os.startfile(nombrepdf)

def printInventario(label1, label2, row):
    texto1=label1.cget("text").split('\n')
    texto2=label2.cget("text").split('\n')
    parrafo = row.cells[0].add_paragraph()
    if label1.cget("bg") == "green":
        font = parrafo.add_run(texto1[0]+" "+texto1[1]).font
        font.color.rgb = RGBColor(11, 170, 35)
        font.bold = True
    elif label1.cget("bg") == "orange":
        font = parrafo.add_run(texto1[0]+" "+texto1[1]).font
        font.color.rgb = RGBColor(175, 83, 0)
    else:
        font = parrafo.add_run(texto1[0]+" "+texto1[1]).font
        font.color.rgb = RGBColor(255, 0, 0)
    parrafo = row.cells[1].add_paragraph()
    if label2.cget("bg") == "green":
        font = parrafo.add_run(texto2[0]+" "+texto2[1]).font
        font.color.rgb = RGBColor(11, 170, 35)
        font.bold = True
    elif label2.cget("bg") == "orange":
        font = parrafo.add_run(texto2[0]+" "+texto2[1]).font
        font.color.rgb = RGBColor(175, 83, 0)
    else:
        font = parrafo.add_run(texto2[0]+" "+texto2[1]).font
        font.color.rgb = RGBColor(255, 0, 0)

ventana = tkinter.Tk()
ventana.title("Examen de Orientación Vocacional")
frame1 = tkinter.Frame()
frame2 = tkinter.Frame()
frame3 = tkinter.Frame()
frame4 = tkinter.Frame()
frame5 = tkinter.Frame()
frame6 = tkinter.Frame()

frame1.pack()
frame2.pack()
frame3.pack()
frame4.pack()
frame5.pack()
frame6.pack()

saludo = tkinter.Label(text="Selecciona el archivo .xlsx con las respuestas:", master=frame1)
direccion = tkinter.Label(text="Direccion Elegida: ", width=70, bg="grey", master=frame1)
botonExaminar = tkinter.Button(text="Examinar...", command=browseFiles, master=frame1)
nombre = tkinter.Label(text="Seleciona el nombre del alumno:", master=frame2)
nombre_alumno = ttk.Combobox(state="readonly", width=40, master=frame2, height = 5)
botonBuscar = tkinter.Button(text="Calcular", command=buscarAlumno, master=frame2)

label_intereses = tkinter.Label(text="\nIntereses del Alumno:", master=frame3)
ss_intereses = tkinter.Label(text="Servicio Social:\n", bg="grey", master=frame3, height = 2, relief = tkinter.SUNKEN)
ep_intereses = tkinter.Label(text="Ejecutivo Persuasiva:\n", bg="grey", master=frame3, height = 2, relief = tkinter.SUNKEN)
v_intereses = tkinter.Label(text="Verbal:\n", bg="grey", master=frame3, height = 2, relief = tkinter.SUNKEN)
ap_intereses = tkinter.Label(text="Artístico Plástico:\n", bg="grey", master=frame3, height = 2, relief = tkinter.SUNKEN)
m_intereses = tkinter.Label(text="Musical:\n", bg="grey", master=frame3, height = 2, relief = tkinter.SUNKEN)
o_intereses = tkinter.Label(text="Organización:\n", bg="grey", master=frame3, height = 2, relief = tkinter.SUNKEN)
cient_intereses = tkinter.Label(text="Científica:\n", bg="grey", master=frame3, height = 2, relief = tkinter.SUNKEN)
calc_intereses = tkinter.Label(text="Cálculo:\n", bg="grey", master=frame3, height = 2, relief = tkinter.SUNKEN)
mc_intereses = tkinter.Label(text="Mecánico Constructiva:\n", bg="grey", master=frame3, height = 2, relief = tkinter.SUNKEN)
al_intereses = tkinter.Label(text="Aire Libre:\n", bg="grey", master=frame3, height = 2, relief = tkinter.SUNKEN)

label_aptitudes = tkinter.Label(text="\nAptitudes del Alumno:", master=frame4)
ss_aptitudes = tkinter.Label(text="Servicio Social:\n", bg="grey", master=frame4, height = 2, relief = tkinter.SUNKEN)
ep_aptitudes = tkinter.Label(text="Ejecutivo Persuasiva:\n", bg="grey", master=frame4, height = 2, relief = tkinter.SUNKEN)
v_aptitudes = tkinter.Label(text="Verbal:\n", bg="grey", master=frame4, height = 2, relief = tkinter.SUNKEN)
ap_aptitudes = tkinter.Label(text="Artístico Plástico:\n", bg="grey", master=frame4, height = 2, relief = tkinter.SUNKEN)
m_aptitudes = tkinter.Label(text="Musical:\n", bg="grey", master=frame4, height = 2, relief = tkinter.SUNKEN)
o_aptitudes = tkinter.Label(text="Organización:\n", bg="grey", master=frame4, height = 2, relief = tkinter.SUNKEN)
cient_aptitudes = tkinter.Label(text="Científica:\n", bg="grey", master=frame4, height = 2, relief = tkinter.SUNKEN)
calc_aptitudes = tkinter.Label(text="Cálculo:\n", bg="grey", master=frame4, height = 2, relief = tkinter.SUNKEN)
mc_aptitudes = tkinter.Label(text="Mecánico Constructiva:\n", bg="grey", master=frame4, height = 2, relief = tkinter.SUNKEN)
dm_aptitudes = tkinter.Label(text="Destreza Manual:\n", bg="grey", master=frame4, height = 2, relief = tkinter.SUNKEN)

grafica = tkinter.Label(master=frame5)
botonPdf = tkinter.Button(text="Generar DOCX", command=generarDocx, master=frame5)

credits = tkinter.Label(text="Programado y diseñado por Luis Fernando Zarza (2020). luis.zarza.progra@gmail.com", font=("Courier", 8), master=frame6)

saludo.pack()
direccion.pack(side = tkinter.LEFT)
botonExaminar.pack(side = tkinter.RIGHT)
nombre.pack(side = tkinter.TOP)
nombre_alumno.pack(side =  tkinter.LEFT)
botonBuscar.pack(side = tkinter.RIGHT)

label_intereses.pack()
ss_intereses.pack(side = tkinter.LEFT)
ep_intereses.pack(side = tkinter.LEFT)
v_intereses.pack(side = tkinter.LEFT)
ap_intereses.pack(side = tkinter.LEFT)
m_intereses.pack(side = tkinter.LEFT)
o_intereses.pack(side = tkinter.LEFT)
cient_intereses.pack(side = tkinter.LEFT)
calc_intereses.pack(side = tkinter.LEFT)
mc_intereses.pack(side = tkinter.LEFT)
al_intereses.pack(side = tkinter.LEFT)

label_aptitudes.pack()
ss_aptitudes.pack(side = tkinter.LEFT)
ep_aptitudes.pack(side = tkinter.LEFT)
v_aptitudes.pack(side = tkinter.LEFT)
ap_aptitudes.pack(side = tkinter.LEFT)
m_aptitudes.pack(side = tkinter.LEFT)
o_aptitudes.pack(side = tkinter.LEFT)
cient_aptitudes.pack(side = tkinter.LEFT)
calc_aptitudes.pack(side = tkinter.LEFT)
mc_aptitudes.pack(side = tkinter.LEFT)
dm_aptitudes.pack(side = tkinter.LEFT)

grafica.pack()
botonPdf.pack()

credits.pack(side = tkinter.RIGHT)

ventana.mainloop()
